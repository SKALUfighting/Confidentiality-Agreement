import streamlit as st
import pandas as pd
from docx import Document
import os
from datetime import datetime
import io
import re
import requests
from urllib.parse import quote

# -------------------- 页面配置 --------------------
st.set_page_config(
    page_title="保密协议生成器 | 国联新创",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="expanded"
)

# -------------------- 自定义CSS样式 --------------------
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 0.5rem;
        padding-top: 1rem;
    }
    .sub-header {
        text-align: center;
        color: #64748B;
        margin-bottom: 2rem;
        font-size: 1.1rem;
    }
    .stButton>button {
        background-color: #3B82F6;
        color: white;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #2563EB;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.2);
    }
    .success-box {
        padding: 1.5rem;
        border-radius: 0.5rem;
        background-color: #D1FAE5;
        border: 1px solid #10B981;
        margin: 1.5rem 0;
    }
    .info-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #EFF6FF;
        border: 1px solid #3B82F6;
        margin: 1rem 0;
    }
    .warning-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #FEF3C7;
        border: 1px solid #F59E0B;
        margin: 1rem 0;
    }
    .step-box {
        background-color: #F8FAFC;
        border-left: 4px solid #3B82F6;
        padding: 1rem;
        margin-bottom: 1.5rem;
        border-radius: 0 0.5rem 0.5rem 0;
    }
    .company-card {
        background: white;
        border: 1px solid #E2E8F0;
        border-radius: 0.5rem;
        padding: 1.25rem;
        margin-bottom: 1rem;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
    }
</style>
""", unsafe_allow_html=True)


# -------------------- 核心函数定义 --------------------

def load_template():
    """加载并验证Word模板文件"""
    template_path = "保密协议模板.docx"

    if not os.path.exists(template_path):
        st.error(f"❌ **关键错误**：未找到模板文件 '{template_path}'")
        st.info("""
        **解决方法：**
        1. 请将您的《保密协议模板.docx》文件放在与此程序相同的目录下
        2. 确保模板中包含以下精确的占位符文本：
           - `[千寻智能(杭州)科技有限公司]`
           - `[浙江省杭州市萧山区宁围街道利一路188号天人大厦浙大研究院数字经济孵化器4层401室-38]`
        """)
        st.stop()

    # 验证模板中是否包含必要的占位符
    try:
        doc = Document(template_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        required_placeholders = [
            "[千寻智能(杭州)科技有限公司]",
            "[浙江省杭州市萧山区宁围街道利一路188号天人大厦浙大研究院数字经济孵化器4层401室-38]"
        ]

        missing = []
        for placeholder in required_placeholders:
            if placeholder not in full_text:
                missing.append(placeholder)

        if missing:
            st.error(f"❌ **模板验证失败**：模板中缺少以下占位符：")
            for m in missing:
                st.code(m, language="text")
            st.info("请在模板文件中添加上述占位符，然后重启应用。")
            st.stop()

        return template_path
    except Exception as e:
        st.error(f"读取模板文件时出错：{str(e)}")
        st.stop()


def smart_replace_in_document(doc, replace_pairs):
    """
    智能替换文档中的文本（增强版）
    处理跨多个Run的文本替换问题
    """
    # 1. 替换所有段落
    for para in doc.paragraphs:
        original_text = para.text
        new_text = original_text

        # 对每个占位符进行替换
        for old, new in replace_pairs.items():
            if old in new_text:
                new_text = new_text.replace(old, new)

        # 如果文本发生了变化，更新段落
        if new_text != original_text:
            # 清空所有runs
            for run in para.runs:
                run.text = ""
            # 重新设置文本到第一个run
            if para.runs:
                para.runs[0].text = new_text
            else:
                # 如果没有run，添加一个
                para.add_run(new_text)

    # 2. 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original_text = para.text
                    new_text = original_text

                    for old, new in replace_pairs.items():
                        if old in new_text:
                            new_text = new_text.replace(old, new)

                    if new_text != original_text:
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = new_text

    return doc


def search_company_address_amap(company_name):
    """
    使用高德地图API搜索公司地址
    """
    # 高德地图API Key（请替换为你的实际API Key）
    api_key = "9f98bd4b65721b9ff59a3b05d1ef0a0d"  # 请替换为实际的高德API Key

    if api_key == "YOUR_AMAP_API_KEY":
        st.warning("⚠️ 请配置高德地图API Key")
        return None

    try:
        # 使用POI搜索API
        url = "https://restapi.amap.com/v3/place/text"
        params = {
            "keywords": company_name,
            "types": "商务写字楼|产业园区|企业|公司",
            "city": "全国",
            "citylimit": "false",
            "output": "json",
            "key": api_key,
            "offset": 10,
            "page": 1,
            "extensions": "base"
        }

        response = requests.get(url, params=params, timeout=10)

        if response.status_code == 200:
            data = response.json()

            if data.get("status") == "1" and int(data.get("count", 0)) > 0 and data.get("pois"):
                poi = data["pois"][0]
                # 提取地址信息
                address = poi.get("address", "")
                pname = poi.get("pname", "")  # 省名
                cityname = poi.get("cityname", "")  # 市名
                adname = poi.get("adname", "")  # 区县名
                name = poi.get("name", "")  # POI名称

                # 组合完整地址
                if address:
                    # 如果有详细地址，组合省市区和详细地址
                    full_address = f"{pname}{cityname}{adname}{address}"
                else:
                    # 如果没有详细地址，使用POI名称作为地址
                    full_address = f"{pname}{cityname}{adname}{name}"

                if full_address.strip():
                    return full_address.strip()

        # 备选方案：使用输入提示API
        tips_url = "https://restapi.amap.com/v3/assistant/inputtips"
        tips_params = {
            "keywords": company_name,
            "type": "商务写字楼|产业园区|企业|公司",
            "city": "全国",
            "output": "json",
            "key": api_key
        }

        tips_response = requests.get(tips_url, params=tips_params, timeout=10)
        if tips_response.status_code == 200:
            tips_data = tips_response.json()
            if tips_data.get("status") == "1" and int(tips_data.get("count", 0)) > 0 and tips_data.get("tips"):
                tip = tips_data["tips"][0]
                tip_address = tip.get("address", "")
                tip_name = tip.get("name", "")

                if tip_address:
                    return tip_address
                elif tip_name:
                    return tip_name

        return None

    except Exception as e:
        st.error(f"高德地图API调用出错: {str(e)}")
        return None


def safe_filename(text, max_length=50):
    """
    生成安全的文件名，移除不安全的字符
    """
    safe_text = re.sub(r'[^\w\s()（）\-]', '', text)
    safe_text = re.sub(r'\s+', '_', safe_text).strip()
    return safe_text[:max_length]


def generate_document(company_name, company_address, template_path):
    """生成新的保密协议文档"""
    doc = Document(template_path)
    replace_pairs = {
        "[千寻智能(杭州)科技有限公司]": company_name,
        "[浙江省杭州市萧山区宁围街道利一路188号天人大厦浙大研究院数字经济孵化器4层401室-38]": company_address
    }
    doc = smart_replace_in_document(doc, replace_pairs)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# -------------------- 主应用界面 --------------------

def main():
    st.markdown('<h1 class="main-header">📄 保密协议智能生成器</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">国联新创 · 内部工具 · created by 宋佳璐</p>', unsafe_allow_html=True)

    template_path = load_template()

    # 初始化session state
    if "company_address" not in st.session_state:
        st.session_state.company_address = ""
    if "search_performed" not in st.session_state:
        st.session_state.search_performed = False
    if "document_ready" not in st.session_state:
        st.session_state.document_ready = False

    col1, col2 = st.columns([2, 1])

    with col1:
        # st.markdown("### 填写协议信息")

        with st.container():
            st.markdown('<div class="step-box"><strong>步骤1：输入公司全称</strong></div>', unsafe_allow_html=True)

            company_name = st.text_input(
                "输入完按Enter键 ",
                # placeholder="输入完按Enter键",
                # help="请务必确保公司名称准确无误",
                key="company_name"
            )

            # 自动搜索地址功能
            if company_name and company_name != st.session_state.get("last_company_name", ""):
                st.session_state.last_company_name = company_name
                st.session_state.search_performed = False
                st.session_state.document_ready = False

                with st.spinner("正在自动搜索公司地址..."):
                    company_address = search_company_address_amap(company_name)

                    if company_address:
                        st.session_state.company_address = company_address
                        st.session_state.search_performed = True
                        st.success(f"✅ 地址已获取: {company_address}")
                    else:
                        st.session_state.company_address = ""
                        st.warning("未找到地址，请手动填写")
                        st.session_state.search_performed = True

        with st.container():
            st.markdown('<div class="step-box"><strong>步骤2：确认公司地址</strong></div>', unsafe_allow_html=True)

            # 地址输入区域
            if st.session_state.search_performed:
                if st.session_state.company_address:
                    # 有自动搜索结果
                    company_address = st.text_area(
                        "请确认或修改，按Ctrl+Enter键",
                        value=st.session_state.company_address,
                        # placeholder="请确认或修改公司注册地址，按Ctrl+Enter键",
                        height=100,
                        key="address_input"
                    )
                else:
                    # 需要手动输入
                    company_address = st.text_area(
                        "公司注册地址 *",
                        # placeholder="请准确填写公司的工商注册地址",
                        height=100,
                        key="address_input"
                    )

                # 检查地址是否填写完成
                if company_address and company_address.strip():
                    st.session_state.company_address = company_address
                    st.session_state.document_ready = True
                    st.success("✅ 地址已确认")
                else:
                    st.session_state.document_ready = False
            else:
                company_address = ""
                st.info("👆 请输入公司名称以触发自动搜索")

        # 步骤3：下载按钮
        if st.session_state.document_ready and company_address and company_address.strip():
            st.markdown('<div class="step-box"><strong>步骤3：下载保密协议</strong></div>', unsafe_allow_html=True)

            # 生成文件名
            current_date = datetime.now().strftime("%Y%m%d")
            safe_name = safe_filename(company_name, 50)
            download_name = f"保密协议_{safe_name}_{current_date}.docx"

            # 直接下载按钮 - 点击后直接下载文件
            # if st.button("📥 下载保密协议", type="primary", use_container_width=True):
            #     with st.spinner("正在生成并下载文档..."):
            # 直接生成文档并提供下载
            file_stream = generate_document(company_name, company_address, template_path)

            # 直接触发下载
            st.download_button(
                label="下载保密协议",
                data=file_stream,
                file_name=download_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )

    with col2:
        st.markdown("### 📖 使用说明")

        with st.expander("操作流程", expanded=True):
            st.markdown("""
            1. **输入公司全称**
               - 系统自动搜索地址

            2. **确认地址信息**
               - 检查并确认地址，也可手动填写

            3. **一键下载**
               - 点击下载按钮直接获取文件
            """)

        # st.divider()

        # st.markdown("### ⚙️ 系统状态")
        # try:
        #     # 模板检查
        #     doc = Document(template_path)
        #     file_size = os.path.getsize(template_path) / 1024
        #     st.success(f"✅ 模板正常 ({file_size:.1f} KB)")
        #
        #     # 占位符检查
        #     full_text = "\n".join([para.text for para in doc.paragraphs[:5]])
        #     placeholders = [
        #         "[千寻智能(杭州)科技有限公司]",
        #         "[浙江省杭州市萧山区宁围街道利一路188号天人大厦浙大研究院数字经济孵化器4层401室-38]"
        #     ]
        #     found_all = all(p in full_text for p in placeholders)
        #     if found_all:
        #         st.success("✅ 占位符就绪")
        #     else:
        #         st.warning("⚠️ 占位符缺失")
        #
        # except Exception as e:
        #     st.error(f"❌ 模板异常: {str(e)}")
        #
        # st.divider()

        # # API配置说明
        # st.markdown("### 🔧 API配置")
        # st.info("""
        # **高德地图API配置：**
        # 1. 注册高德开放平台账号
        # 2. 创建应用获取API Key
        # 3. 将API Key填入代码中的 `YOUR_AMAP_API_KEY` 位置
        # 4. 确保API服务已开通：POI搜索、输入提示
        # """)


if __name__ == "__main__":
    # 初始化session state
    session_keys = ["company_name", "company_address", "search_performed",
                    "last_company_name", "document_ready"]
    for key in session_keys:
        if key not in st.session_state:
            st.session_state[key] = ""

    main()
