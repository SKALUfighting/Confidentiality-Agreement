from docx import Document

doc = Document('保密协议模板.docx')
target_text = '[千寻智能(杭州)科技有限公司]' # 要查找的文本

found = False
for i, para in enumerate(doc.paragraphs):
    if target_text in para.text:
        print(f'✅ 在第 {i+1} 段中找到目标文本：“{para.text.strip()}”')
        found = True

if not found:
    print('❌ 未在段落中找到目标文本，请检查：')
    print('1. 模板文件名是否正确？')
    print('2. 占位符括号是英文[]吗？有多余空格吗？')
    print('--- 前3段内容预览 ---')
    for i, para in enumerate(doc.paragraphs[:3]):
        print(f'段{i+1}: {para.text}')